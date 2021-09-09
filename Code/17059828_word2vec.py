import nltk
from nltk.corpus import brown
from nltk.corpus import stopwords
from nltk.corpus import reuters
import string
import gensim
import os
from nltk.tokenize import word_tokenize, sent_tokenize
import pandas as pd
import random
import pickle
import glob
import numpy as np
import sys
from docx import Document
import gensim.downloader as api
import time


# Word2Vec - text8 - v1

# timer 
tic1 = time.perf_counter()

temp3 = api.load('text8')
corp2 = temp3

# sg = type (0 = CBOW, 1 = Skip-Gram), min_count = words of count less than stated are ignored
model2 = gensim.models.Word2Vec(corp2, size = 180, min_count = 2, sg = 0, iter = 10) 
model2.save("Text8Model.model")

# once done training - delete model an store as keyedVector to save RAM - wv = word vector
word_vectors2 = model2.wv
del model2

# time
toc1 = time.perf_counter()
print("Text8 model gen time; ", toc1 - tic1)

# word2vec tests on model to check functionality
testing = word_vectors2.most_similar(positive = ['low'])
test2 = word_vectors2.similarity('explode', 'currency')
print(testing, "\n", "\n", test2, "\n")
print(word_vectors2['aircraft'])




# word2vec - brown + reuters + risk assessments - V2
#fpath = r'C:\Users\bowde\Desktop\MSc Project\RiskAssessmentDocs\TestData/' # used for V2 (obsolete - inaccurate)
# punct = ['!','"','#','$','%','&',"'",'(',')','*','+',',','-','.','/',':',';','<','=','>','?','@','[','\\',']','^','_','`','{','|','}','~','``',"''",'--']

# train_corp = []
# _docx = []
# _docx_forToken = []
# def Train_doc_open(file):
#     _doc = Document(file)
#     table = _doc.tables[0]
#     _docx.append(_doc.paragraphs[0].text)
#     for i, row in enumerate(table.rows):
#         text = (cell.text for cell in row.cells)
#         if i == 0: # set key - header text
#             keys = tuple(text)
#             continue
#         row_data = dict(zip(keys, text))
#         x = list(row_data.items())
#         _docx_forToken.extend(x)
    
# for file in glob.glob(os.path.join(fpath, '*.docx')): 
#     Train_doc_open(file)
#     train_corp.append(file)

# print("training corpus: ", train_corp)
# print("\n", "\n")
 
# train_token = []
# for x in _docx_forToken:
#     if x[0]:
#         train_token.extend(list(word_tokenize(x[1])))
#         train_token = [word for word in train_token if word != "."]
#         train_token = [word for word in train_token if word != ","]
#         train_token = [x.lower() for x in train_token]
                       
# print("\n", "\n")
# # print(train_token)
# print("\n", "\n")

# # timer
# tic2 = time.perf_counter()

# corp = []
# temp = brown.words()
# temp2 = reuters.words()
# temp += train_token
# temp += temp2
# _corp = [x.lower() for x in temp]
# to_remove = punct + stopwords.words('english')
# _corp = [x for x in temp if x not in to_remove]
# random.shuffle(_corp)
# for i in _corp:
#     corp.extend(gensim.utils.simple_preprocess(i))
# print(len(corp), "\n")

# # window = max dist between target word, workers = num of threads, sg = type (0 = CBOW, 1 = Skip-Gram)
# model = gensim.models.Word2Vec(sentences = [corp], size = 300, window = 4, min_count = 3, workers = 4, sg = 0, iter = 240) 
# model.save("TestModel.model")
# model = gensim.models.Word2Vec.load("TestModel.model")

# # once done training - delete model an store as keyedVector to save RAM - wv = word vector
# word_vectors = model.wv
# del model

# # time
# toc2 = time.perf_counter()
# print("Brown + Reuters + UserDocs model gen time; ", toc2 - tic2)

# vect = word_vectors['dog'] # get vector of word
# sims = word_vectors.most_similar(positive = ['woman', 'king'], negative = ['man'], topn = 8)
# comp = word_vectors.similarity(w1 = 'low', w2 = 'lower') # cosine
# print(vect, "\n")
# print(sims, "\n")
# print(comp, "\n")



# Word2Vec - pretrained model

# google news trained file contains ~= 3 million words, 300 dimentions - google pre-trained model - https://drive.google.com/file/d/0B7XkCwpI5KDYNlNUTTlSS21pQmM/edit?resourcekey=0-wjGZdNAUop6WykTtMip30g
# file = 'GoogleNews-vectors-negative300.bin.gz'
# model = gensim.models.KeyedVectors.load_word2vec_format(file, binary = True)
# # print(model['low'], '\n')
# print(model.most_similar(positive = ['Low', 'low'], negative = ['High', 'high']), '\n')

