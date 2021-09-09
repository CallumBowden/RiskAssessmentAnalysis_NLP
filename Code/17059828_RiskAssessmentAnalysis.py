# ********** Improvements **********
# fix fpath not working when user changes it at runtime - feature removed for now
# add some more charts, code logs etc. to show efficiency of program and word2vec
# display data mixing doc titles - fix
# scroll bars - QoL
# user feedback for no items found - blacklist, scores
# fine-tune rating gen

import gensim as g
import os
from nltk.tokenize import word_tokenize, sent_tokenize
import glob
from docx import Document
import nltk
from nltk.corpus import brown
from nltk.corpus import reuters
import pandas as pd
import tkinter as tk
from functools import partial
import psutil
import difflib as dif
from nltk.tokenize import PunktSentenceTokenizer
import time
import numpy as np
import sys

# ****** IMPORTANT ******
# Generate word2vec model via the 17059828_word2vec.py
# change file directory via line 31

fpath = r'C:\Users\bowde\Desktop\MSc Project\RiskAssessmentDocs\TestData/' # chnage file directory here

tmr_startup = time.perf_counter()
doc_cont = [] # list of file path
_docx = [] # dict array
_docx_forToken = [] # dict array to convert to list
model = g.models.Word2Vec.load("Text8Model.model") # word2vec model
wordVec = model.wv
train_punkt = [] # sent list to train punkt from
tmr_punkt_addCorp = time.perf_counter()
for w in brown.sents(): # about 20-30 seconds
    train_punkt.append(w)
# for w2 in reuters.sents(): # about 50-60 seconds with brown corpus
#     train_punkt.append(w2)
tmr_punkt_addCorp_End = time.perf_counter()
tmr_punkt_addCorp_total = round(tmr_punkt_addCorp_End - tmr_punkt_addCorp, 3)

# events below
def read_table(doc, table_num = 1, nheader = 1):
    table = doc.tables[table_num - 1] # grab table
    df = [[cell.text for cell in row.cells] for row in table.rows] # create list of each row
    data = pd.DataFrame(df) # create dataframe
    data = data.rename(columns = data.iloc[0]).drop(data.index[0]).reset_index(drop = True) # set headers
    return data 

def open_all_files():
    def clear_var(): # clear variables to stop duplicates and inflation
        doc_cont.clear()
        _docx.clear()
        _docx_forToken.clear()
    def doc_open(file): # append to relevent variables
        _doc = Document(file)
        _docx.append(_doc.paragraphs[0].text)
        table = _doc.tables[0] # get table in doc
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            if i == 0: # set key - header text
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            _docx.append(row_data)
            x = list(row_data.items())
            _docx_forToken.extend(x)
    def set_open(): # update user
        lst_openf = tk.Listbox(master = form_buttons, height = 2, width = 72, bd = 4)
        lst_openf.insert(1, "Opened; ", doc_cont)
        lst_openf.grid(row = 2, column = 3, sticky = 'nsew', padx = 8, pady = 8)
    def open_loop(): # open files
        for file in glob.glob(os.path.join(fpath, '*.docx')): 
            doc_open(file)
            doc_cont.append(file)
    clear_var()
    open_loop()
    set_open()
    print(doc_cont)
    print('\n')
 
def single_file():
    file1 = tk.StringVar()
    def clearVar(): # clear variables to stop duplicates and inflation
        doc_cont.clear()
        _docx.clear()
        _docx_forToken.clear()
    def file_checker(filename): # check file exists
        clearVar()
        temp = filename.get()
        try:
            for f in glob.glob(os.path.join(fpath, '*.docx')):
                tempFile = os.path.join(fpath, temp)
                if(os.path.samefile(tempFile, f)):
                    open_doc(tempFile)
                    doc_cont.append(os.path.join(fpath, temp))
                    print(doc_cont)
                    print('\n')
                    setOpen()
                    return
        except:
            print("No file with that name")
            clearVar()
    def open_doc(file): # add data to relevent variables
        print("opened")
        _doc = Document(file)
        _docx.append(_doc.paragraphs[0].text)
        table = _doc.tables[0]
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)
            if i == 0: # set key - header text
                keys = tuple(text)
                continue
            row_data = dict(zip(keys, text))
            _docx.append(row_data)
            x = list(row_data.items())
            _docx_forToken.extend(x)
    def setOpen(): # update user
        lst_openf = tk.Listbox(master = form_buttons, height = 2, width = 72, bd = 4)
        lst_openf.insert(1, "Opened; ", doc_cont)
        lst_openf.grid(row = 2, column = 3, sticky = 'nsew', padx = 8, pady = 8)
    def file_entry_window(): # open input box
        win = tk.Toplevel(MainWindow)
        win.title("File Open")
        lbl_entry = tk.Label(master = win, text = "Enter File Name (don't forget .docx)")
        ent_toReturn = tk.Entry(master = win, width = 60, textvariable = file1)
        btn_submit = tk.Button(master = win, text = "Submit", command = partial(file_checker, file1))
        btn_close = tk.Button(master = win, text = "Close", command = win.destroy)
        lbl_entry.pack()
        ent_toReturn.pack()
        btn_submit.pack()
        btn_close.pack()
    file_entry_window()
    print("\n")
    
# def setFP(): # doc_cont being annoying + fpath won't update - fix
#     filePath = tk.StringVar()
#     def winInput(): # open input box
#         win = tk.Toplevel(MainWindow)
#         win.title("Set File Directory")
#         lbl_entry = tk.Label(master = win, text = "Enter a file path (Make Sure it ends with /)")
#         lbl_eg = tk.Label(master = win, text = r"E.g - 'c:\User\Desktop\FilesHere/'")
#         ent_toReturn = tk.Entry(master = win, width = 60, textvariable = filePath)
#         btn_submit = tk.Button(master = win, text = "Submit", command = partial(set_fp, filePath))
#         btn_close = tk.Button(master = win, text = "Close", command = win.destroy)
#         lbl_entry.pack()
#         lbl_eg.pack()
#         ent_toReturn.pack()
#         btn_submit.pack()
#         btn_close.pack()
#     def set_fp(temp): # set new directory
#         fpath = os.path.join(temp.get())
#         print(fpath)
#     winInput()
#     print('\n')
    
def displayStuff(): # list details for viewing
    def display_window(): # display doc contents
        winDisplay = tk.Toplevel(MainWindow)
        winDisplay.title("File Content")
        lbl_data = tk.Listbox(master = winDisplay, height = 50, width = 300, bd = 4)
        for i in range(len(_docx)):
            lbl_data.insert(i, _docx[i])
        lbl_data.grid(row = 0, column = 0, sticky = 'nsew')
    display_window()
    print('\n')
            
def show_scores():
    preScores = []
    postScores = []
    auditorScore = []
    auditorPostScore = []
    flagged_rows = []
    low_prob = [] # list of descriptives for 'low'
    med_prob = [] # list of descriptives for 'medium'
    high_prob = [] # list of descriptives for 'high'
    sev_prob = [] # list of descriptives for 'severe'
    def rating_gen(): # word similarity for scoring - word2vec
        # gets words similar to below, creates list for each category, to allow user to enter similar words.
        _low = wordVec.most_similar('low', topn = 30)
        _low2 = wordVec.most_similar('small', topn = 30) 
        _low3 = wordVec.most_similar('little', topn = 30)
        _low4 = wordVec.most_similar('insignificant', topn = 30)
        for i in _low:
            seq = dif.SequenceMatcher(None, 'low', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                low_prob.append(i[0])
        for i in _low2:
            seq = dif.SequenceMatcher(None, 'small', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                low_prob.append(i[0])   
        for i in _low3:
            seq = dif.SequenceMatcher(None, 'little', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                low_prob.append(i[0]) 
        for i in _low4:
            seq = dif.SequenceMatcher(None, 'insignificant', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                low_prob.append(i[0]) 
        low_prob.append('low') 
        low_prob.append('small') 
        low_prob.append('little') 
        low_prob.append('insignificant') 
        _med = wordVec.most_similar('medium', topn = 30)
        _med2 = wordVec.most_similar('average', topn = 30)
        _med3 = wordVec.most_similar('normal', topn = 30)
        _med4 = wordVec.most_similar('moderate', topn = 30)
        _med5 = wordVec.most_similar('ok', topn = 30)
        for i in _med:
            seq = dif.SequenceMatcher(None, 'medium', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                med_prob.append(i[0])
        for i in _med2:
            seq = dif.SequenceMatcher(None, 'average', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                med_prob.append(i[0])   
        for i in _med3:
            seq = dif.SequenceMatcher(None, 'normal', i[0]) # not producing useful info
            temp = seq.ratio() * 100
            if(temp >= 56):
                med_prob.append(i[0]) 
        for i in _med4:
            seq = dif.SequenceMatcher(None, 'moderate', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                med_prob.append(i[0]) 
        for i in _med5:
            seq = dif.SequenceMatcher(None, 'ok', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                med_prob.append(i[0]) 
        med_prob.append('medium') 
        med_prob.append('average') 
        med_prob.append('normal') 
        med_prob.append('moderate') 
        med_prob.append('ok') 
        _high = wordVec.most_similar('high', topn = 30)
        _high2 = wordVec.most_similar('probable', topn = 30)
        _high3 = wordVec.most_similar('dangerous', topn = 30)
        for i in _high:
            seq = dif.SequenceMatcher(None, 'high', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                high_prob.append(i[0])
        for i in _high2:
            seq = dif.SequenceMatcher(None, 'probable', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                high_prob.append(i[0])   
        for i in _high3:
            seq = dif.SequenceMatcher(None, 'dangerous', i[0])
            temp = seq.ratio() * 100
            if(temp >= 56):
                high_prob.append(i[0]) 
        high_prob.append('high') 
        high_prob.append('probable')
        high_prob.append('dangerous') 
        _sev = wordVec.most_similar('severe', topn = 30)
        _sev2 = wordVec.most_similar('certain', topn = 30)
        _sev3 = wordVec.most_similar('lethal', topn = 30)
        _sev4 = wordVec.most_similar('deadly', topn = 30)
        _sev5 = wordVec.most_similar('guaranteed', topn = 30)
        for i in _sev:
            seq = dif.SequenceMatcher(None, 'severe', i[0]) # not producing useful info
            temp = seq.ratio() * 100
            if(temp >= 56):
                sev_prob.append(i[0])
        for i in _sev2:
            seq = dif.SequenceMatcher(None, 'certain', i[0]) # not producing useful info
            temp = seq.ratio() * 100
            if(temp >= 56):
                sev_prob.append(i[0])
        for i in _sev3:
            seq = dif.SequenceMatcher(None, 'lethal', i[0]) # not producing useful info
            temp = seq.ratio() * 100
            if(temp >= 56):
                sev_prob.append(i[0])
        for i in _sev4:
            seq = dif.SequenceMatcher(None, 'deadly', i[0]) # not producing useful info
            temp = seq.ratio() * 100
            if(temp >= 56):
                sev_prob.append(i[0])
        for i in _sev5:
            seq = dif.SequenceMatcher(None, 'guaranteed', i[0]) # not producing useful info
            temp = seq.ratio() * 100
            if(temp >= 56):
                sev_prob.append(i[0])
        sev_prob.append('severe') 
        sev_prob.append('certain') 
        sev_prob.append('lethal') 
        sev_prob.append('deadly') 
        sev_prob.append('very-high') 
        sev_prob.append('guaranteed')
    def display_scores(): # display calculated scores
        winDisplay = tk.Toplevel(MainWindow)
        winDisplay.title("Scores")
        lbl_data = tk.Listbox(master = winDisplay, height = 40, width = 24, bd = 4)
        lbl_postData = tk.Listbox(master = winDisplay, height = 40, width = 24, bd = 4)
        lbl_auditorScore = tk.Listbox(master = winDisplay, height = 40, width = 24, bd = 4)
        lbl_auditorPostScore = tk.Listbox(master = winDisplay, height = 40, width = 24, bd = 4)
        for x in range(len(preScores)):
            lbl_data.insert(0, "Risk Score; ", preScores[x])
        for x in range(len(preScores)):
            lbl_postData.insert(0, "Post-Risk Score; ", postScores[x])
        for x in range(len(auditorScore)):
            lbl_auditorScore.insert(0, "Auditor Score; ", auditorScore[x])
        for x in range(len(auditorPostScore)):
            lbl_auditorPostScore.insert(0, "Auditor Post-Risk Score; ", auditorPostScore[x])
        lbl_data.grid(row = 0, column = 0, sticky = 'nsew')
        lbl_postData.grid(row = 0, column = 1, sticky = 'nsew')
        lbl_auditorScore.grid(row = 0, column = 2, sticky = 'nsew')
        lbl_auditorPostScore.grid(row = 0, column = 3, sticky = 'nsew')
    def calc_scores(): # calculate scores using _data dataframe
        # column titles are rigid at the moment - improve in future
        for i in range(len(_data)):
            try:
                x = 0
                y = 0
                z = 0
                human_asmt = _data.loc[i, "Risk_Factor"] # auditors opinion of total risk
                human_asmt = human_asmt.lower()
                for low1 in range(len(low_prob)):
                    l = low_prob[low1]
                    if (l == human_asmt):
                        x = 1
                        continue
                for med1 in range(len(med_prob)):
                    m = med_prob[med1]
                    if (m in human_asmt):
                        x = 2
                        continue
                for high1 in range(len(high_prob)):
                    h = high_prob[high1]
                    if (h in human_asmt):
                        x = 3
                        continue
                for sev1 in range(len(sev_prob)): 
                    s = sev_prob[sev1]
                    if (s in human_asmt):
                        x = 4
                        continue
                risk_prob = _data.loc[i, "Risk_Probability"] # auditors opinion of risk probability
                risk_prob = risk_prob.lower()
                for low2 in range(len(low_prob)):
                    l = low_prob[low2]
                    if (l == risk_prob):
                        y = 1
                for med2 in range(len(med_prob)):
                    m = med_prob[med2]
                    if (m in risk_prob):
                        y = 2
                for high2 in range(len(high_prob)):
                    h = high_prob[high2]
                    if (h in risk_prob):
                        y = 3
                for sev2 in range(len(sev_prob)): 
                    s = sev_prob[sev2]
                    if (s in risk_prob):
                        y = 4
                risk_sev = _data.loc[i, "Risk_Severity"]  # auditors opinion of risk severity
                risk_sev = risk_sev.lower()
                for low3 in range(len(low_prob)):
                    l = low_prob[low3]
                    if (l == risk_sev):
                        z = 1
                for med3 in range(len(med_prob)):
                    m = med_prob[med3]
                    if (m in risk_sev):
                        z = 2
                for high3 in range(len(high_prob)):
                    h = high_prob[high3]
                    if (h in risk_sev):
                        z = 3
                for sev3 in range(len(sev_prob)): 
                    s = sev_prob[sev3]
                    if (s in risk_sev):
                        z = 4
                try: # mitigations
                    a = 0
                    b = 0
                    c = 0
                    post_asmt = _data.loc[i, "Mitigation_Risk_Factor"] # auditors opinion of total risk after changes
                    post_asmt = post_asmt.lower()
                    for low4 in range(len(low_prob)):
                        l = low_prob[low4]
                        if (l == post_asmt):
                            a = 1
                    for med4 in range(len(med_prob)):
                        m = med_prob[med4]
                        if (m in post_asmt):
                            a = 2
                    for high4 in range(len(high_prob)):
                        h = high_prob[high4]
                        if (h in post_asmt):
                            a = 3
                    for sev4 in range(len(sev_prob)): 
                        s = sev_prob[sev4]
                        if (s in post_asmt):
                            a = 4
                    post_risk_prob = _data.loc[i, "Mitigation_Probability"] # auditors opinion of risk probability after changes
                    post_risk_prob = post_risk_prob.lower()
                    for low5 in range(len(low_prob)):
                        l = low_prob[low5]
                        if (l == post_risk_prob):
                            b = 1
                    for med5 in range(len(med_prob)):
                        m = med_prob[med5]
                        if (m in post_risk_prob):
                            b = 2
                    for high5 in range(len(high_prob)):
                        h = high_prob[high5]
                        if (h in post_risk_prob):
                            b = 3
                    for sev5 in range(len(sev_prob)): 
                        s = sev_prob[sev5]
                        if (s in post_risk_prob):
                            b = 4
                    post_risk_sev = _data.loc[i, "Mitigation_Severity"]  # auditors opinion of risk severity after changes
                    post_risk_sev = post_risk_sev.lower()
                    for low6 in range(len(low_prob)):
                        l = low_prob[low6]
                        if (l == post_risk_sev):
                            c = 1
                    for med6 in range(len(med_prob)):
                        m = med_prob[med6]
                        if (m in post_risk_sev):
                            c = 2
                    for high6 in range(len(high_prob)):
                        h = high_prob[high6]
                        if (h in post_risk_sev):
                            c = 3
                    for sev6 in range(len(sev_prob)): 
                        s = sev_prob[sev6]
                        if (s in post_risk_sev):
                            c = 4
                except: # for if no mitigations exist
                    print("Mitigation Pass")
                calc = (x * z) / (5 - y)
                post_calc = (a * c) / (5 - b)
                if (post_calc >= calc and post_calc != 0):
                    flagged_rows.append(_data.loc[i].values)
                else:
                    if (calc >= x + 1):
                        flagged_rows.append(_data.loc[i].values)
                    else:
                        if (post_calc >= a + 1):
                            flagged_rows.append(_data.loc[i].values)
                preScores.append(calc)
                postScores.append(post_calc)
                auditorScore.append(x)
                auditorPostScore.append(a)
            except: # if no risk rating exists
                print("Risk Pass")
    def display_flagged(): # display rows of concern
        i = 0
        winDataframe = tk.Toplevel(MainWindow)
        winDataframe.title("Flagged Rows")
        lbl_rows = tk.Listbox(master = winDataframe, height = 40, width = 280, bd = 4)
        for i in range(len(flagged_rows)):
            temp = flagged_rows[i]    
            lbl_rows.insert(i, temp)
            i + 1
        lbl_rows.grid(row = 0, column = 0, sticky = 'nsew')  
    rating_gen()
    for fp in range(len(doc_cont)):
        doc = Document(doc_cont[fp]) 
        _data = read_table(doc, 1, 1)
        calc_scores()
    display_scores()
    display_flagged()
    print('\n')
    
def blacklist():
    obj_list = []
    bl_items = []
    bl_matched = []
    flaggedRows = []
    def tokenize(): # tokenize snetance array to words
        train_token = [] # tokenized list
        for x in _docx_forToken:
            if x[0]:
                word = x[1].lower()
                train_token.extend(list(word_tokenize(word)))
                train_token = [word for word in train_token if word != "."]
                train_token = [word for word in train_token if word != ","]
        return train_token
    def nameEntRec(lst_tokenized): # apply pos tags and add tools etc. to list
        def process(sent): # apply pos
            sent = nltk.pos_tag(sent)
            return sent
        temp = process(lst_tokenized)
        for i in temp:
            if (i[1] == 'NN'): # Noun
                obj_list.append(i[0])
            if (i[1] == 'NNS'): # Noun plural
                obj_list.append(i[0])
            if (i[1] == 'NNP'): # proper noun
                obj_list.append(i[0])
            if (i[1] == 'NNPS'): # proper noun plural
                obj_list.append(i[0])
    def blacklistItems(): # find items that match blacklist
        bl_file = open("Blacklisted_Items.txt", "r")
        lines = bl_file.readlines()
        for line in enumerate(lines):
            temp = line[1].lower()
            temp = temp.strip()
            bl_items.append(temp)
        for word in bl_items:
            if (word in obj_list):
                bl_matched.append(word)
        bl_file.close()
    def displayBLItems(): # display items that are flagged
        winDisplay = tk.Toplevel(MainWindow)
        winDisplay.title("Blacklisted Items")
        lst_data = tk.Listbox(master = winDisplay, height = 40, width = 60, bd = 4)
        lbl_colLabel = tk.Label(master = winDisplay, text =  "Items that matched Blacklisted Items in opened document(s);")
        for i in range(len(bl_matched)):
            lst_data.insert(0, bl_matched[i])
        lbl_colLabel.grid(row = 0, column = 0, sticky = 'nsew')
        lst_data.grid(row = 1, column = 0, sticky = 'nsew')
    def getFlaggedEntry(): # find rows containing flagged rows
        for i in range(len(_data)):
            flag = False # prevent duplicates
            for c in range(len(_data.columns)):
                if (flag == True):
                    continue
                else:
                    for j in bl_matched:
                            try:
                                if (j in _data.iloc[i, c]):
                                    flaggedRows.append(_data.loc[i].values)
                                    flag = True
                                    continue
                            except:
                                continue
    def displayFlaggedEntry(): # display flagged rows containing flagged items
        i = 0
        winDataframe = tk.Toplevel(MainWindow)
        winDataframe.title("Flagged Rows")
        lst_rows = tk.Listbox(master = winDataframe, height = 40, width = 280, bd = 4)
        for i in range(len(flaggedRows)):
            temp = flaggedRows[i]
            lst_rows.insert(i, temp)
        lst_rows.grid(row = 0, column = 0, sticky = 'nsew')     
    temp = tokenize()
    nameEntRec(temp)
    blacklistItems()
    displayBLItems()
    for fp in range(len(doc_cont)):
        doc = Document(doc_cont[fp]) 
        _data = read_table(doc, 1, 1)
        getFlaggedEntry()
    displayFlaggedEntry()
    print('\n')

def POS_List():
    nouns = []
    plural_nouns = []
    proper_nouns = []
    proper_plural_nouns = []
    tagged = []
    def tokenize(): # punkt_sent_tokenize with used corpora as training
        tmr_punkt_train = time.perf_counter()
        sent_tokenizer = PunktSentenceTokenizer(" ".join(map(str, train_punkt)))
        tmr_punkt_train_end = time.perf_counter()
        tmr_punkt_train_total = round(tmr_punkt_train_end - tmr_punkt_train, 3)
        tmr_punkt_total = round(tmr_punkt_addCorp_total + tmr_punkt_train_total, 3)
        print("Punkt Train time = ", tmr_punkt_train_total)
        print("Total time Punkt uses = ", tmr_punkt_total)
        for x in _docx_forToken: # POS tags - list of tuples
            temp = x[1]
            x_tokenized = sent_tokenizer.tokenize(temp)
            try:
                for i in x_tokenized:
                    word = nltk.word_tokenize(i)
                    tagged.append(nltk.pos_tag(word))
            except:
                print("Error with list")
        return tagged       
    def nameEntRec(lst_tokenized): # find nouns (tools, names, objects etc.)
        for i in range(len(lst_tokenized)):
            for j in lst_tokenized[i]:
                if (j[1] == 'NN'): # noun
                    nouns.append(j[0])
                if (j[1] == 'NNS'): # noun plural
                    plural_nouns.append(j[0])
                if (j[1] == 'NNP'): # proper noun
                    proper_nouns.append(j[0])
                if (j[1] == 'NNPS'): # proper noun plural
                    proper_plural_nouns.append(j[0])
    def display_POS(): # display results
        winDisplay = tk.Toplevel(MainWindow)
        winDisplay.title("POS Filter")
        lst_n = tk.Listbox(master = winDisplay, height = 50, width = 22, bd = 4)
        lst_np = tk.Listbox(master = winDisplay, height = 50, width = 22, bd = 4)
        lst_pn = tk.Listbox(master = winDisplay, height = 50, width = 22, bd = 4)
        lst_pnp = tk.Listbox(master = winDisplay, height = 50, width = 22, bd = 4)
        lbl_col4 = tk.Label(master = winDisplay, text = "Proper Noun (Plural); ")
        lbl_col3 = tk.Label(master = winDisplay, text = "Proper Noun; ")
        lbl_col2 = tk.Label(master = winDisplay, text = "Noun (Plural); ")
        lbl_col1 = tk.Label(master = winDisplay, text = "Noun; ")
        for x in range(len(nouns)):
            lst_n.insert(0, nouns[x])
        for x in range(len(plural_nouns)):
            lst_np.insert(0, plural_nouns[x])
        for x in range(len(proper_nouns)):
            lst_pn.insert(0, proper_nouns[x])
        for x in range(len(proper_plural_nouns)):
            lst_pnp.insert(0, proper_plural_nouns[x])
        lbl_col1.grid(row = 0, column = 0, sticky = 'nsew')
        lbl_col2.grid(row = 0, column = 1, sticky = 'nsew')
        lbl_col3.grid(row = 0, column = 2, sticky = 'nsew')
        lbl_col4.grid(row = 0, column = 3, sticky = 'nsew')
        lst_n.grid(row = 1, column = 0, sticky = 'nsew')
        lst_np.grid(row = 1, column = 1, sticky = 'nsew')
        lst_pn.grid(row = 1, column = 2, sticky = 'nsew')
        lst_pnp.grid(row = 1, column = 3, sticky = 'nsew')
    temp = tokenize()
    nameEntRec(temp)
    display_POS()
    print("\n")
    
# main window
MainWindow = tk.Tk()
MainWindow.title("Risk Assessment Analysis - Navigation Popup")
# frame
form_buttons = tk.Frame()
form_buttons.pack(fill = tk.Y, ipadx = 10, ipady = 10)

# buttons/widgets/info etc.
btn_open = tk.Button(master = form_buttons, text = "Open A File", command = single_file)
btn_openAll = tk.Button(master = form_buttons, text = "Open All Files", command = open_all_files)
btn_analysis = tk.Button(master = form_buttons, text = "Display Scores", command = show_scores)
btn_blacklist = tk.Button(master = form_buttons, text = "Display Blacklist", command = blacklist)
btn_display = tk.Button(master = form_buttons, text = "Display Opened Content", command = displayStuff)
btn_close = tk.Button(master = form_buttons, text = "Close", command = MainWindow.destroy)
btn_pos = tk.Button(master = form_buttons, text = "POS Filter", command = POS_List)
# btn_Setfp = tk.Button(master = form_buttons, text = "Set File Path", command = setFP)
lst_fp = tk.Listbox(master = form_buttons, height = 2, width = 72, bd = 4)
lst_fopen = tk.Listbox(master = form_buttons, height = 2, width = 72, bd = 4)

# layout formatting
lst_fp.insert(0, "Current path; ", fpath)
lst_fopen.insert(0, "Opened; ", doc_cont)
# btn_Setfp.grid(row = 0, column = 0, sticky = "nsew", padx = 8, pady = 8)
lst_fp.grid(row = 1, column = 3, sticky = 'nsew', padx = 8, pady = 8)
lst_fopen.grid(row = 2, column = 3, sticky = 'nsew', padx = 8, pady = 8)
btn_open.grid(row = 1, column = 0, sticky = "nsew", padx = 8, pady = 8)
btn_openAll.grid(row = 2, column = 0, sticky = "nsew", padx = 8, pady = 8)
btn_analysis.grid(row = 1, column = 1, sticky= "nsew", padx = 8, pady = 8)
btn_blacklist.grid(row = 2, column = 1, sticky= "nsew", padx = 8, pady = 8)
btn_display.grid(row = 1, column = 2, sticky = "nsew", padx = 8, pady = 8)
btn_close.grid(row = 0, column = 3, sticky = "nsew", padx = 8, pady = 8)
btn_pos.grid(row = 2, column = 2, sticky = "nsew", padx = 8, pady = 8)

# misc labels
label = tk.Label(text = "Word2Vec Model Used - Train in IDE")
label2 = tk.Label(text = "POS uses a corpus to train - Change in IDE - TAKES TIME TO LOAD (Program Might Freeze For A Second)")
label.pack()
label2.pack()

# stats on open
print("Used CPU; " , psutil.cpu_percent(), "%")
mem = psutil.virtual_memory()
print("Used RAM; ", mem.percent, "%")
print("\n")

# program loop
tmr_startup_end = time.perf_counter()
print("Startup = ", round(tmr_startup_end - tmr_startup, 3))
MainWindow.mainloop()

# stats on close
print("Used CPU; " , psutil.cpu_percent(), "%")
mem2 = psutil.virtual_memory()
print("Used RAM; ", mem2.percent, "%")
print("\n")

